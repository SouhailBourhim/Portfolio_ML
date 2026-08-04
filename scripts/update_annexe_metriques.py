"""Remove the retracted equivalence claim from the evaluation-metrics annex.

Addresses: P4 — this annex described the block bootstrap as the tool that
"autorise la conclusion « statistiquement indiscernables »". It does not. The
bootstrap produces MARGINAL intervals, one per strategy; overlapping marginal
intervals are not a test of the difference between two strategies, and failing
to reject is not accepting the null. The annex was therefore teaching the
wrong inference in the very document that explains the project's evaluation
metrics — the worst possible place for it.

What the bootstrap genuinely does is prevent a point ranking from being
presented as a fact. What tests a difference is the PAIRED bootstrap on the
return differences, which is a different instrument and is named as such.

Targeted edits: this file has no generator, and the replacement text belongs
in reviewable source rather than only inside a binary. Idempotent — each
replacement is skipped once the new wording is present.

NOTE for whoever next revises this annex: it documents the block bootstrap
(§4.2) but has no section on the PAIRED bootstrap, which is the instrument
that actually tests a difference and the basis of every comparative claim
the project now makes. The text below therefore points at the code rather
than at a section number. Adding that section is worthwhile and is left as
a deliberate follow-up, not done here.
"""

from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "Annexe_Metriques_Evaluation.docx"
REALITY_CHECK = ROOT / "data" / "gold" / "reality_check_results.json"

REPLACEMENTS = (
    (
        "et c'est lui qui autorise la conclusion « statistiquement "
        "indiscernables » au lieu d'un classement ponctuel présenté comme un fait",
        "et c'est lui qui empêche de présenter un classement ponctuel comme un "
        "fait. Il n'autorise pas pour autant la conclusion inverse : des "
        "intervalles MARGINAUX qui se chevauchent ne testent pas la différence "
        "entre deux stratégies. Tester cette différence exige un bootstrap "
        "PAIRÉ, mené sur les écarts de rendement eux-mêmes — instrument absent "
        "de cette annexe et implémenté dans metrics.paired_block_bootstrap, "
        "résultats dans data/gold/paired_comparison_results.json"
    ),
    (
        "et ce sont elles qui permettent d'affirmer « statistiquement "
        "indiscernable » avec des preuves, au lieu de publier une estimation "
        "ponctuelle en espérant qu'elle tienne",
        "et ce sont elles qui permettent de dire ce qui est établi et ce qui ne "
        "l'est pas, au lieu de publier une estimation ponctuelle en espérant "
        "qu'elle tienne. La formulation licite est « aucune surperformance "
        "n'est établie », jamais « les stratégies sont équivalentes » : une "
        "équivalence demanderait un test dédié contre une marge fixée à "
        "l'avance, qui n'a pas été mené"
    ),
)


def set_paragraph(paragraph, text: str) -> None:
    """Rewrite a paragraph, keeping run[0] as the style carrier."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    keeper = runs[0]
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    keeper.text = text


def set_if_different(paragraph, text: str) -> bool:
    """Set visible text only when the generated wording actually changed."""
    if paragraph.text == text:
        return False
    set_paragraph(paragraph, text)
    return True


def _paragraph_starting(document: Document, prefix: str):
    """Return the unique paragraph whose visible text begins with ``prefix``."""
    matches = [p for p in document.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(
            f"Expected exactly one paragraph starting {prefix!r}, found {len(matches)}."
        )
    return matches[0]


def _insert_after(document: Document, anchor, text: str, template):
    """Insert text after ``anchor``, retaining the adjacent paragraph's formatting."""
    paragraph = document.add_paragraph()
    if template._p.pPr is not None:
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    set_paragraph(paragraph, text)
    if template.runs and template.runs[0]._r.rPr is not None:
        paragraph.runs[0]._r.get_or_add_rPr()
        paragraph.runs[0]._r.rPr.getparent().replace(
            paragraph.runs[0]._r.rPr, deepcopy(template.runs[0]._r.rPr)
        )
    anchor._p.addnext(paragraph._p)
    return paragraph


def _reality_check_summary() -> tuple[int, int, int]:
    """Read released candidate and SPA-retention counts; never type results into Word."""
    raw = json.loads(REALITY_CHECK.read_text(encoding="utf-8"))
    universes = raw["universes"].values()
    candidates = {u["n_candidates"] for u in universes}
    retained = [
        row["spa_candidates_retained"]
        for universe in raw["universes"].values()
        for row in universe["tests"].values()
    ]
    if len(candidates) != 1:
        raise ValueError(f"Expected one candidate-set size, found {sorted(candidates)}.")
    return candidates.pop(), min(retained), max(retained)


def _insert_phase5_sections(document: Document) -> int:
    """Add the released paired and multiple-testing protocol once, in place."""
    if any(p.text.startswith("4.3 Bootstrap pairé") for p in document.paragraphs):
        return 0

    n_candidates, retained_min, retained_max = _reality_check_summary()
    old_heading = _paragraph_starting(document, "4.3 Registre des essais")
    old_heading.text = "4.5 Registre des essais (DSRTrialLedger)"

    # The final paragraph of §4.2 must point to the section that now exists.
    bootstrap_note = _paragraph_starting(document, "C'est l'outil d'honnêteté")
    set_paragraph(
        bootstrap_note,
        "C'est l'outil d'honnêteté le plus important du projet : il transforme "
        "« 1,12 contre 0,97 » en deux intervalles marginaux et empêche de "
        "présenter un classement ponctuel comme un fait. Il n'autorise pas pour "
        "autant la conclusion inverse : des intervalles MARGINAUX qui se "
        "chevauchent ne testent pas la différence entre deux stratégies. Cette "
        "différence est traitée par le bootstrap PAIRÉ de la section 4.3."
    )

    ledger_text = _paragraph_starting(document, "Le DSR ne déflate correctement")
    set_paragraph(
        ledger_text,
        "Le registre rend la recherche inspectable : pour la Phase 5, il consigne "
        "51 essais hiérarchiques (15 configurations de signal évaluées par IC et "
        "36 essais de leviers avec séries de rendements). Ce registre rend visible "
        "la sélection, mais il ne remplace pas la correction de recherche de la "
        "section 4.4 : celle-ci évalue les 240 chemins de portefeuille atteignables "
        "sur la même fenêtre de test gelée."
    )

    heading_template = old_heading
    body_template = ledger_text
    anchor = bootstrap_note
    heading = _insert_after(document, anchor, "4.3 Bootstrap pairé sur les différences", heading_template)
    body = _insert_after(
        document,
        heading,
        "Lorsqu'il faut comparer deux stratégies, les deux séries de rendements "
        "nets sont ré-échantillonnées avec les MÊMES blocs circulaires de 21 jours. "
        "Chaque tirage conserve donc la dépendance temporelle de chaque stratégie "
        "et leur corrélation le même jour. L'objet testé est directement l'écart de "
        "rendement et l'écart de Sharpe, pas le chevauchement de deux intervalles "
        "marginaux séparés.",
        body_template,
    )
    paired = _insert_after(
        document,
        body,
        "Le résultat publié comprend l'intervalle pairé, une p-value unilatérale "
        "centrée sous l'hypothèse nulle d'absence de surperformance et la "
        "probabilité descriptive que l'écart de Sharpe soit positif. Une "
        "surperformance n'est établie que si l'intervalle exclut zéro et si la "
        "p-value respecte le seuil déclaré. Ne pas rejeter l'hypothèse nulle ne "
        "démontre jamais l'équivalence : cela exigerait un test dédié contre une "
        "marge définie à l'avance.",
        body_template,
    )
    heading = _insert_after(document, paired, "4.4 Correction du choix parmi de nombreuses variantes", heading_template)
    body = _insert_after(
        document,
        heading,
        f"White Reality Check et Hansen SPA testent l'hypothèse composite qu'aucun "
        f"des {n_candidates} candidats atteignables ne surperforme un benchmark sur "
        "la fenêtre de test gelée. Toutes les variantes sont évaluées sur les mêmes "
        "dates et partagent les mêmes tirages bootstrap : les variantes proches ne "
        "sont pas artificiellement traitées comme des paris indépendants.",
        body_template,
    )
    body = _insert_after(
        document,
        body,
        f"Le benchmark primaire est fixé avant lecture des p-values : "
        "regime_conditional, le système que la couche ML devait améliorer. "
        "equal_weight est exploratoire : le battre n'établit pas une valeur ajoutée "
        "ML, car le système de régimes et le max_sharpe classique le dépassent déjà. "
        "White RC et Hansen SPA sont toujours rapportés ensemble. SPA a retenu "
        f"{retained_min} à {retained_max} candidats sur {n_candidates} ; l'écart "
        "avec White vient donc principalement de la studentisation, non d'un "
        "élagage massif. Huit comparaisons externes sont rapportées (2 univers, 2 "
        "benchmarks, 2 statistiques) sans correction supplémentaire : leurs "
        "résultats exploratoires ne doivent pas être sélectionnés a posteriori.",
        body_template,
    )
    return 4


def _update_traceability_table(document: Document) -> int:
    """List the two Phase-5 controls in the code-location table exactly once."""
    table = document.tables[2]
    labels = [row.cells[0].text for row in table.rows]
    additions = (
        ("Bootstrap pairé", "P4", "metrics.py : paired_block_bootstrap"),
        ("White RC + Hansen SPA", "P4", "reality_check.py : evaluate_candidate_set"),
    )
    changed = 0
    for label, problem, implementation in additions:
        if label in labels:
            continue
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = label, problem, implementation
        changed += 1
    # The expanded map is one logical table. Slightly tightening its type keeps
    # its final row and Word's required trailing paragraph on the same page,
    # rather than producing an empty last page.
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", "30"), ("bottom", "30"),
                                ("start", "80"), ("end", "80")):
                margin = tc_mar.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), value)
                margin.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return changed


def _polish_section_four_and_pagination(document: Document) -> int:
    """Keep the expanded P4 section accurate and avoid a stranded final paragraph."""
    overview = _paragraph_starting(
        document, "C'est la partie qui distingue ce projet d'un backtest ordinaire"
    )
    changed = int(set_if_different(
        overview,
        "C'est la partie qui distingue ce projet d'un backtest ordinaire. Ces cinq "
        "contrôles ne disent pas si la performance est bonne : ils disent si elle "
        "est CRÉDIBLE.",
    ))

    conclusion = _paragraph_starting(document, "Lecture d'ensemble.")
    changed += int(set_if_different(
        conclusion,
        "Lecture d'ensemble. Les sections 4 et 7 séparent une observation "
        "historique d'une conclusion établie : ni des intervalles marginaux ni "
        "l'absence de rejet ne prouvent l'équivalence."
    ))
    traceability = _paragraph_starting(document, "8. Traçabilité P1–P4")
    # The two added P4 sections make the final synthesis spill onto a nearly
    # blank page. It is a synthesis of the metrics, so placing it immediately
    # before the implementation map preserves the reading logic and lets the
    # map occupy the final page as a coherent unit.
    if conclusion._p.getnext() is not traceability._p:
        traceability._p.addprevious(conclusion._p)
        changed += 1
    # A Word table is followed by one legacy blank paragraph in this source
    # file. Once the table gained two rows, that otherwise-empty paragraph was
    # alone on a fifth page. It carries no formatting or content, so remove it.
    last = document.paragraphs[-1]
    if not last.text and not last.runs:
        last._element.getparent().remove(last._element)
        changed += 1
    return changed


def main() -> None:
    document = Document(str(DOCX))
    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text
        updated = text
        for old, new in REPLACEMENTS:
            if old in updated:
                updated = updated.replace(old, new)
        if updated != text:
            set_paragraph(paragraph, updated)
            changed += 1
    changed += _insert_phase5_sections(document)
    changed += _update_traceability_table(document)
    changed += _polish_section_four_and_pagination(document)
    if changed:
        document.save(str(DOCX))
    print(f"{DOCX.relative_to(ROOT)}: {changed} paragraph(s) rewritten")


if __name__ == "__main__":
    main()
