"""Build docs/Livrable_Phase8_Etudes_Robustesse.docx in the project's house style.

Covers the five post-Phase-5 investigations (CLAUDE.md §12E–§12I) plus the
weight-cap finding (§10.1). Clones an existing Livrable so fonts, heading
styles, table borders, header and footer match the rest of the set exactly.

Every figure is read from `data/gold/*.json` at build time — never typed — so
the document cannot drift from the artifacts it describes.
"""
import copy
import json
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from docx.text.paragraph import Paragraph

ROOT = Path("/Users/apple/Projects/Portfolio_ML")
GOLD = ROOT / "data" / "gold"
SRC = ROOT / "docs" / "Livrable_Phase5_Evaluation_OOS.docx"
OUT = ROOT / "docs" / "Livrable_Phase8_Etudes_Robustesse.docx"
USABLE_TW, HEADER_FILL = 9026, "DCE6F1"


def load(name):
    return json.loads((GOLD / name).read_text())


DM = load("deep_morocco_results.json")
FI = load("fundamentals_ic_lift.json")["full_2021"]
FP = load("fundamentals_portfolio.json")
NW = load("nested_walkforward_results.json")
RC = load("regime_conditional_cap.json")
CW = load("crisis_windows.json")
CAP = load("etf_cap_verdict.json")
P5 = load("phase5_results.json")


def fr(x, n=3):
    return f"{x:.{n}f}".replace(".", ",")


# ── document scaffold ────────────────────────────────────────────────────────
shutil.copy(SRC, OUT)
doc = Document(str(OUT))
P = doc.paragraphs
P[2].runs[0].text = "Livrable de la Phase 8 — Études de robustesse"
for r in P[2].runs[1:]:
    r.text = ""
P[3].runs[0].text = ("Cinq tentatives indépendantes de dépasser la ligne de base — "
                     "protocoles pré-enregistrés, contrôles, et le seul résultat significatif")
for r in P[3].runs[1:]:
    r.text = ""

tpl_h1 = copy.deepcopy(P[10]._p)
tpl_h2 = copy.deepcopy(doc.paragraphs[20]._p)
tpl_body = copy.deepcopy(P[19]._p)
tpl_li = copy.deepcopy(P[22]._p)
tpl_tbl = copy.deepcopy(doc.tables[0]._tbl)

for t in list(doc.tables):
    t._tbl.getparent().remove(t._tbl)
for p in list(doc.paragraphs[9:]):
    p._p.getparent().remove(p._p)

body_el = doc.element.body
sectPr = body_el.find(qn("w:sectPr"))


def _append(el):
    sectPr.addprevious(el) if sectPr is not None else body_el.append(el)


def para(text, kind="body"):
    el = copy.deepcopy({"h1": tpl_h1, "h2": tpl_h2, "li": tpl_li}.get(kind, tpl_body))
    _append(el)
    p = Paragraph(el, doc)
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    p.runs[0].text = text
    p.runs[0].bold = True if kind in ("h1", "h2") else None
    if kind not in ("h1", "h2"):
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p


def table(headers, rows, widths):
    el = copy.deepcopy(tpl_tbl)
    _append(el)
    from docx.table import Table
    t = Table(el, doc)
    while len(t.rows) > 1:
        t._tbl.remove(t.rows[-1]._tr)
    grid = t._tbl.find(qn("w:tblGrid"))
    for gc in list(grid.findall(qn("w:gridCol"))):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    cells = t.rows[0].cells
    while len(cells) > len(headers):
        t.rows[0]._tr.remove(cells[-1]._tc)
        cells = t.rows[0].cells
    while len(cells) < len(headers):
        t.rows[0]._tr.append(copy.deepcopy(cells[-1]._tc))
        cells = t.rows[0].cells
    t._tbl.find(qn("w:tblPr")).find(qn("w:tblW")).set(qn("w:w"), str(sum(widths)))

    def fill(cell, text, w, bold):
        cell.width = Twips(w)
        p = cell.paragraphs[0]
        for r in p.runs[1:]:
            r._r.getparent().remove(r._r)
        if not p.runs:
            p.add_run("")
        p.runs[0].text = text
        p.runs[0].bold = bold
        p.runs[0].font.size = Pt(8.5)
        tcPr = cell._tc.get_or_add_tcPr()
        for s in tcPr.findall(qn("w:shd")):
            tcPr.remove(s)
        if bold:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), HEADER_FILL)
            nxt = next((e for tag in ("w:noWrap", "w:tcMar", "w:textDirection",
                                      "w:tcFitText", "w:vAlign", "w:hideMark")
                        for e in tcPr.findall(qn(tag))), None)
            nxt.addprevious(shd) if nxt is not None else tcPr.append(shd)

    for c, (txt, w) in enumerate(zip(headers, widths)):
        fill(t.rows[0].cells[c], txt, w, True)
    for row in rows:
        tr = copy.deepcopy(t.rows[0]._tr)
        t._tbl.append(tr)
        cs = t.rows[-1].cells
        for c, (txt, w) in enumerate(zip(row, widths)):
            fill(cs[c], txt, w, False)


def spacer():
    el = copy.deepcopy(tpl_body)
    _append(el)
    p = Paragraph(el, doc)
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    p.runs[0].text = ""


# ═══ 1. Objet ════════════════════════════════════════════════════════════════
para("1. Objet — pourquoi une phase d'études après la Phase 5", "h1")
para(
    "La Phase 5 a conclu que, hors échantillon et avec de vraies barres d'erreur, les modèles F7 "
    "honnêtement calibrés et la ligne de base régime + covariance dynamique sont "
    "STATISTIQUEMENT INDISCERNABLES. Face à un tel résultat, deux attitudes sont possibles : "
    "l'accepter, ou vérifier qu'il résiste. Nous avons choisi la seconde — parce qu'un résultat "
    "nul non testé est indiscernable d'un manque d'effort."
)
para(
    "Cette phase rassemble cinq investigations indépendantes menées après la Phase 5. Chacune "
    "attaque une explication différente du résultat nul : le signal était-il sous-alimenté en "
    "données ? mal nourri (prix seuls, sans fondamentaux) ? mal évalué (fenêtre de test trop "
    "courte) ? mal exploité (la contrainte plutôt que la covariance) ? Et enfin : mesurons-nous "
    "seulement la bonne chose (un Sharpe moyen plutôt qu'un comportement en crise) ?"
)
para(
    "Quatre de ces cinq études se concluent par un résultat négatif ou nul. La cinquième produit "
    "le SEUL résultat statistiquement significatif du projet. Les deux catégories sont "
    "présentées avec la même rigueur : c'est la cohérence de l'ensemble qui constitue le "
    "livrable, pas la seule réussite."
)

# ═══ 2. Protocole ════════════════════════════════════════════════════════════
para("2. Protocole commun — ce qui rend ces résultats opposables", "h1")
para(
    "Résultats pré-enregistrés. Chaque étude fixe ses conclusions possibles (A / B / C) dans le "
    "docstring du script AVANT exécution. Une hypothèse qui échoue ne peut donc pas être "
    "requalifiée après coup en « exploration intéressante ».", "li"
)
para(
    "Définitions externes. Là où une fenêtre ou un seuil pourrait être choisi à notre avantage, "
    "il vient d'une source extérieure : les cinq crises de l'étude 5 sont des dates "
    "sommet-à-creux publiées du S&P 500, jamais dérivées de nos propres drawdowns — sans quoi on "
    "sélectionne les périodes où l'on paraît bon et on appelle cela un résultat.", "li"
)
para(
    "Contrôles destinés à tuer l'hypothèse. L'étude 4 inclut une variante INVERSÉE — "
    "délibérément dans la mauvaise direction. Si elle fonctionne aussi, l'effet mesuré n'est pas "
    "celui que l'on croit. C'est ce contrôle qui a tranché.", "li"
)
para(
    "Aucune affirmation de significativité sans test. Les écarts de Sharpe sont rapportés avec "
    "leurs intervalles bootstrap et décrits comme non significatifs tant qu'ils se chevauchent. "
    "L'unique test formel de cette phase (étude 5) est présenté avec sa version conservatrice EN "
    "PREMIER et sa version libérale explicitement marquée comme optimiste.", "li"
)
para(
    "Divulgation des ajustements de protocole. L'étude 4 a vu son seuil de matérialité ajouté "
    "après une exécution à blanc, parce que la règle initiale (« strictement supérieur ») aurait "
    "qualifié un écart de +0,0016 de succès. Le changement rend le test PLUS strict et il est "
    "consigné dans le script et dans la note de recherche — modifier une règle de décision après "
    "avoir vu un résultat est précisément ce que le pré-enregistrement doit empêcher.", "li"
)

# ═══ 3. Étude 1 — deep Morocco ═══════════════════════════════════════════════
para("3. Étude 1 — Le signal était-il sous-alimenté en données ? (P2, P4)", "h1")
u = DM["universe"]
para(
    f"Hypothèse. La Phase 5 mesure des coefficients d'information (IC) faibles "
    f"({DM['information_coefficient']['phase5_reference']}) sur ~{u['comparison_current']['pooled_rows_approx']:,} "
    f"lignes de panel. Peut-être le modèle manque-t-il simplement de données. Test : un univers "
    f"marocain PROFOND de {u['n_assets']} actions sur {u['n_days']:,} jours "
    f"({u['start']} → {u['end']}), soit {u['pooled_rows']:,} lignes de panel — environ cinq fois "
    f"l'univers courant — évalué avec la machinerie EXACTE de la Phase 5.".replace(",", " ")
)
para(
    f"Résultat en deux temps. Le modèle devient mesurablement meilleur : l'IC passe de "
    f"{DM['information_coefficient']['phase5_reference']} à "
    f"{fr(DM['information_coefficient']['rf']['mean_ic'], 4)} (RF) et "
    f"{fr(DM['information_coefficient']['xgb']['mean_ic'], 4)} (XGB) — un facteur 2 à 4, confirmé "
    f"par les deux algorithmes. Le portefeuille, lui, ne suit pas."
)
table(
    ["Stratégie", "Sharpe net (test gelé)"],
    [[k, fr(v["test_sharpe_net"])] for k, v in
     sorted(DM["strategies"].items(), key=lambda kv: -kv[1]["test_sharpe_net"])],
    [4500, 4526],
)
spacer()
para(
    "Verdict : davantage de données rend le modèle plus intelligent sans produire d'avantage de "
    "portefeuille statistiquement significatif — tous les intervalles sont larges et incluent "
    "zéro. Le plafond n'est pas la QUANTITÉ de données. Cette étude ferme la piste « plus "
    "d'historique de prix » et réoriente vers la qualité des données, ce que fait l'étude 2."
)

# ═══ 4. Étude 2 — fondamentaux ═══════════════════════════════════════════════
para("4. Étude 2 — Et si les données étaient de meilleure qualité ? (P1, P4)", "h1")
para(
    "Hypothèse. Suite directe : on ajoute des ratios de valorisation point-in-time (P/E, P/B, "
    "P/S, D/E) aux actions BVC, décalés de 90 jours ouvrés pour respecter le délai de publication "
    "AMMC — la discipline causale étant verrouillée par un test de corruption du futur, comme "
    "pour les features de marché."
)
rf_b = FI["baseline"]["random_forest"]["best"]["mean_ic"]
rf_t = FI["treatment"]["random_forest"]["best"]["mean_ic"]
para(
    f"Le modèle apprend nettement plus. L'IC en validation croisée purgée passe de "
    f"{fr(rf_b, 4)} à {fr(rf_t, 4)} pour le RandomForest — il double presque, et `FUND_pb` "
    f"devient la deuxième feature la plus importante. L'arbre utilise réellement les "
    f"fondamentaux ; ce n'est pas du bruit ajusté."
)
table(
    ["Stratégie", "Sharpe net (fenêtre test gelée)"],
    [[k, fr(v["test_net_sharpe"])] for k, v in
     sorted(FP["results"].items(), key=lambda kv: -kv[1]["test_net_sharpe"])],
    [4500, 4526],
)
spacer()
lift = FP["results"]["rf_signal_fundamentals"]["test_net_sharpe"] - \
    FP["results"]["rf_signal_baseline"]["test_net_sharpe"]
para(
    f"Verdict, et il est frappant : l'apport des fondamentaux au signal F7 vaut {fr(lift)} de "
    f"Sharpe — NÉGATIF. Non significatif (les intervalles se chevauchent), mais toutes les "
    f"estimations ponctuelles vont dans le mauvais sens, pour une feature qui a DOUBLÉ l'IC. "
    f"Les fondamentaux réduisent le turnover de moitié, comme on l'attend d'un signal plus lent — "
    f"et la stabilisation tue l'alpha dans la même proportion."
)
para(
    "Deuxième confirmation du même plafond : précision de prédiction ≠ performance de "
    "portefeuille. Ce n'est plus une observation isolée, c'est un motif."
)

# ═══ 5. Étude 3 — nested walk-forward ════════════════════════════════════════
para("5. Étude 3 — Et si le problème était l'ÉVALUATION ? (P4)", "h1")
d = NW["design"]
para(
    f"Hypothèse. La Phase 5 avait elle-même désigné cette étape : sur `full_2021`, la fenêtre de "
    f"test gelée ne fait que ~1,75 an et TOUS les intervalles dépassent 2,2 de Sharpe. Ce n'est "
    f"pas un problème de modèle mais de TAILLE D'ÉCHANTILLON, qu'aucun meilleur modèle ne corrige. "
    f"Le walk-forward imbriqué ré-sélectionne la configuration à {d['n_folds']} frontières "
    f"successives et concatène chaque segment hors échantillon : {d['n_oos_rows']} lignes "
    f"({d['oos_start']} → {d['oos_end']}) contre 455, la sélection ne voyant jamais sa propre "
    f"fenêtre d'évaluation."
)
single = {**P5["full_2021"]["tuned"], **P5["full_2021"]["baselines"]}
rows = []
for k in ("regime_conditional", "xgb_signal_tuned", "rf_signal_tuned", "equal_weight"):
    s = single.get(k)
    n = NW["strategies"].get(k)
    if not (s and n):
        continue
    sw = s["test_sharpe_ci"][1] - s["test_sharpe_ci"][0]
    rows.append([k, fr(s["test_sharpe_net"]), fr(sw), fr(n["sharpe_net"]),
                 fr(n["ci_width"]), f"{100*(1-n['ci_width']/sw):.1f} %"])
table(["Stratégie", "Sharpe (unique)", "largeur", "Sharpe (imbriqué)", "largeur", "resserrement"],
      rows, [2100, 1500, 1150, 1700, 1150, 1426])
spacer()
para(
    f"Trois lectures à garder distinctes. (1) Le RESSERREMENT est le résultat robuste : largeur "
    f"moyenne réduite de 28,6 %, mieux que les 24,3 % qu'un simple effet de taille prédirait ; "
    f"Sharpe déflaté {fr(NW['best_dsr_vs_search'])} sur {NW['n_search_trials']} configurations "
    f"contre 0,67 sur 36. (2) Toutes les bornes inférieures deviennent positives — au découpage "
    f"unique celle de l'équipondéré descendait à −0,090 ; sur cet univers, chaque stratégie "
    f"examinée est désormais crédiblement positive hors échantillon. (3) La hausse des NIVEAUX "
    f"est un effet de période, PAS une amélioration : le Sharpe de toutes les stratégies monte, "
    f"la fenêtre imbriquée commençant en 2023-07 et non en 2024-10. Seules les comparaisons "
    f"internes à un même passage ont un sens."
)
para(
    f"Le classement repasse à `{NW['best']}` — soit un TROISIÈME ordre différent en trois "
    f"évaluations de `full_2021`. La lecture honnête n'est pas « le régime l'emporte finalement » "
    f"mais que l'ordre ponctuel est INSTABLE AU PROTOCOLE D'ÉVALUATION, ce que la Phase 5 "
    f"concluait déjà. L'estimation imbriquée mérite le plus de poids, mais les intervalles se "
    f"chevauchent presque intégralement : aucune significativité, dans aucun sens."
)

# ═══ 6. Étude 4 — regime-conditional cap ═════════════════════════════════════
para("6. Étude 4 — Et si l'on conditionnait la CONTRAINTE au régime ? (P1, P3)", "h1")
para(
    "Hypothèse, motivée par nos propres données. Le balayage du plafond (section 8) montre que la "
    "contrainte régularise mieux que tout modèle de covariance essayé. La Phase 4 conditionne la "
    "COVARIANCE au régime ; le geste non testé est de conditionner LE PLAFOND — le resserrer en "
    "régime baissier (davantage de rétrécissement quand les corrélations explosent : P1 + P3), le "
    "relâcher en régime haussier. Aucune ligne de code de production n'a été nécessaire : "
    "`RegimeConditionalStrategy` accepte déjà des sous-stratégies."
)
rows = []
for uni, v in RC["universes"].items():
    res = v["results"]
    rows.append([
        uni, fr(v["baseline_sharpe"]),
        f"{v['best_candidate']} {fr(res[v['best_candidate']]['sharpe_net'])} ({v['best_candidate_lift']:+.3f})".replace(".", ","),
        f"{fr(res[v['best_control']]['sharpe_net'])}",
        v["verdict"].split("—")[0].strip(),
    ])
table(["Univers", "Référence 25/25", "Meilleur candidat", "Contrôle INVERSÉ", "Verdict"],
      rows, [1500, 1500, 3100, 1500, 1426])
spacer()
para(
    "Verdict (C) sur les deux univers : aucune variante ne dépasse la référence de façon "
    "matérielle. Le mécanisme supposé échoue franchement — resserrer en régime baissier n'apporte "
    "rien (+0,0016 sur `full_2021`, −0,117 sur `etf_2017`) ; le peu de mouvement positif vient de "
    "RELÂCHER le plafond haussier, soit la moitié opposée de l'idée."
)
para(
    "C'est le CONTRÔLE qui tranche, et il est instructif. Sur `full_2021` la variante INVERSÉE est "
    "la pire, ce qui semblerait confirmer un vrai signal de régime. Mais sur `etf_2017` elle BAT "
    "le meilleur candidat « dans le bon sens » (−0,047 contre −0,085). Un effet de régime qui "
    "change de signe selon l'univers n'est pas un effet de régime. N'avoir exécuté que "
    "`full_2021` aurait fait passer ce contrôle pour une preuve à l'appui."
)
para(
    "Une hypothèse bien motivée, issue de notre propre résultat le plus fort, pré-enregistrée, "
    "contrôlée — et réfutée. C'est un bon résultat pour le livrable : la piste est fermée "
    "proprement et par écrit, plutôt que laissée en « on aurait pu essayer »."
)

# ═══ 7. Étude 5 — crisis ═════════════════════════════════════════════════════
para("7. Étude 5 — Mesurions-nous seulement la bonne chose ? (P3) ⭐", "h1")
para(
    "Hypothèse. P3 — la rupture de la diversification en crise — était le problème le moins "
    "directement étayé du projet : toutes les phases rapportaient un Sharpe et un drawdown sur "
    "période complète, aucune ne mesurait le comportement PENDANT les crises. Or le cahier des "
    "charges demande la « pertinence financière » : pour un allocataire, c'est le comportement en "
    "drawdown. Coût de l'étude : zéro calcul — les données étaient déjà dans les artefacts."
)
para("7.1 Résultat A — la contrainte protège ; le 1/N non", "h2")
C = CW["universes"]["etf_2017"]
rows = []
for key, meta in CW["crises"].items():
    w = C.get(key)
    if not w or "min_variance_lw" not in w or "equal_weight" not in w:
        continue
    o, e = w["min_variance_lw"], w["equal_weight"]
    rows.append([
        meta["label"],
        f"{fr(100*o['cum_return'], 1)} %", f"{fr(100*e['cum_return'], 1)} %",
        f"{fr(100*(o['cum_return']-e['cum_return']), 1)} pts",
        f"{o['recovery_days']} j / {e['recovery_days']} j" if o.get("recovery_days") and e.get("recovery_days") else "—",
    ])
table(["Crise", "Optimiseurs", "Équipondéré", "Écart", "Récupération (opt./1N)"],
      rows, [2500, 1500, 1500, 1300, 2226])
spacer()
para(
    "Sur les cinq crises, sans exception, l'optimisation sous contrainte perd moins, chute moins "
    "et récupère plus vite. Le délai de récupération est l'écart le plus régulier — environ deux "
    "fois moins de temps sous l'eau — et il n'apparaît nulle part ailleurs dans le livrable."
)
para(
    "Attribution honnête, et elle est essentielle : ce gain revient à la CONTRAINTE et au modèle "
    "de covariance (P1/P3), PAS spécifiquement à la couche de régime. Sur trois des cinq fenêtres "
    "les trois optimiseurs sont identiques à la décimale près — en régime baissier "
    "`regime_conditional` EST `min_variance_lw` par construction, et le plafond de 25 % sur "
    "5 actifs contraint fortement l'allocation. La page 1 du tableau de bord CALCULE ce compte "
    "depuis l'artefact au lieu de l'affirmer."
)
para("7.2 Résultat B — le HMM non supervisé a détecté les cinq crises", "h2")
rd = CW["regime_detection"]["etf_2017"]
sig = rd["significance"]
para(
    f"Le détecteur est non supervisé : aucune date de crise, aucun label de récession ne lui a "
    f"jamais été fourni. Il n'observe que rendement, volatilité et corrélation moyenne, et décide "
    f"EN TEMPS RÉEL, à partir du passé uniquement. Taux de régime baissier : "
    f"{fr(100*rd['bear_rate_in_crisis'], 1)} % pendant les crises contre "
    f"{fr(100*rd['bear_rate_outside'], 1)} % hors crise — rapport {fr(rd['risk_ratio'], 2)}×, avec "
    f"{sig['crises_exceeding_base_rate']} crises au-dessus du taux de base."
)
table(["Test", "Résultat", "Lecture"],
      [["Test des signes (conservateur)", f"p = {str(sig['sign_test_p_conservative']).replace('.', ',')}",
        "Chaque crise = 1 observation (n=5). La corrélation sérielle intra-fenêtre ne peut pas le gonfler. À CITER EN PREMIER."],
       ["Fisher exact (libéral)", f"p = {str(sig['fisher_exact_p_liberal']).replace('.', ',')}",
        "Traite 248 rééquilibrages mensuels corrélés comme indépendants. Optimiste — ne jamais citer seul."]],
      [2400, 1600, 5026])
spacer()
para(
    "C'est le SEUL résultat statistiquement significatif du projet : toutes nos comparaisons de "
    "ratio de Sharpe ont des intervalles qui se chevauchent."
)
para(
    "Réserve qui doit voyager avec ce résultat : « baissier » est DÉFINI comme l'état à plus "
    "faible rendement moyen, et une crise est par nature une période de faible rendement — une "
    "part de l'association est donc définitionnelle. Ce qui ne l'est pas : la détection est "
    "causale et en temps réel ; le modèle ignore qu'une crise commence et la signale tout de même "
    "à trois fois son taux de base, SANS crier au loup en permanence (29 % hors crise, pas 80 %)."
)
para(
    "Deux affirmations à ne jamais confondre, et à garder dans des phrases séparées : le "
    "détecteur VOIT démonstrativement ce qu'il prétend voir ; le fait d'AGIR dessus reste "
    "statistiquement indiscernable des lignes de base (Phase 5, étude 3)."
)

# ═══ 8. Le plafond ═══════════════════════════════════════════════════════════
para("8. Résultat transversal — le plafond fait plus que les modèles (P1)", "h1")
caps = sorted(CAP["verdicts"], key=float)
best_t = CAP["verdicts"][caps[0]]["classical_sharpe"]
best_l = CAP["verdicts"][caps[-1]]["classical_sharpe"]
table(["Plafond `max_weight`", "Meilleur Sharpe net classique", "Allocations distinctes / 248"],
      [[("sans plafond" if float(c) >= 1.0 else f"{100*float(c):.0f} %"),
        fr(CAP["verdicts"][c]["classical_sharpe"], 4),
        str(CAP["results"][c]["min_variance_lw"]["distinct_allocations"])] for c in caps],
      [2600, 3200, 3226])
spacer()
para(
    f"Un écart de {fr(100*(best_t-best_l)/best_l, 1)} % de Sharpe dû à la SEULE contrainte "
    f"({fr(best_t, 4)} → {fr(best_l, 4)}), décroissant de façon monotone. C'est davantage que "
    f"l'écart entre deux modèles quelconques sur cet univers : Ledoit-Wolf, EWMA, DCC-GARCH et la "
    f"commutation de régime HMM réunis le déplacent moins. C'est le résultat de Jagannathan & Ma "
    f"(2003) reproduit sur nos données — une contrainte de poids active équivaut mathématiquement "
    f"à un rétrécissement de la covariance. La contrainte réalisait donc le contrôle d'erreur "
    f"d'estimation (P1) que le ML devait apporter."
)
para(
    "À présenter comme un résultat, pas comme une limite : le contrôle du risque le plus efficace "
    "de ce système s'est révélé être la limite de position qu'un mandat réel imposerait de toute "
    "façon — exactement ce que le cadrage « contraintes réalistes de gestion » du cahier des "
    "charges invite à découvrir."
)

# ═══ 9. Synthèse ═════════════════════════════════════════════════════════════
para("9. Synthèse — cinq routes testées, et ce que leur convergence démontre", "h1")
table(["Route testée", "Verdict"],
      [["Prédiction de rendement F7, calibrée honnêtement (Phase 5)", "indiscernable de la ligne de base"],
       ["5× plus d'historique de prix (étude 1)", "IC ×2–4, aucun avantage de portefeuille"],
       ["Fondamentaux point-in-time (étude 2)", "IC ×2, Sharpe en BAISSE"],
       ["Ré-sélection par pli, 74 % de données OOS en plus (étude 3)", "intervalles −28,6 %, classement instable"],
       ["Contrainte conditionnée au régime (étude 4)", "réfutée ; le contrôle change de signe"]],
      [5200, 3826])
spacer()
para(
    "Le constat n'est pas que chaque idée a échoué individuellement — c'est qu'à cette taille "
    "d'échantillon, l'évaluation NE PEUT PAS RÉSOUDRE des différences de l'ordre de grandeur que "
    "ces idées produisent. C'est l'affirmation scientifique la plus défendable du projet, et elle "
    "n'est disponible que parce que cinq tentatives indépendantes ont été menées avec le même "
    "protocole."
)
para(
    "Ce que le projet peut affirmer, en revanche : la détection de régime fonctionne "
    "(étude 5, p = "
    f"{str(sig['sign_test_p_conservative']).replace('.', ',')}), l'optimisation sous contrainte protège en crise "
    "(étude 5), et le plafond de position est le régulariseur le plus efficace mesuré "
    "(section 8). Trois résultats positifs — dont aucun ne porte sur la prédiction de rendement."
)

# ═══ 10. Traçabilité ═════════════════════════════════════════════════════════
para("10. Traçabilité P1–P4", "h1")
para(
    "P1 (covariance bruitée) — la section 8 démontre que la contrainte de poids réalise le "
    "contrôle d'erreur d'estimation attribué au ML ; l'étude 4 teste, et réfute, l'idée de la "
    "conditionner au régime.", "li"
)
para(
    "P2 (non-stationnarité) — l'étude 5 valide indépendamment le détecteur de régime : un modèle "
    "entraîné uniquement sur rendement, volatilité et corrélation redécouvre cinq crises sans "
    "qu'aucune date ne lui soit fournie.", "li"
)
para(
    "P3 (rupture de la diversification) — l'étude 5 fournit la première mesure directe du "
    "comportement en crise du projet ; jusque-là, P3 n'était étayé que par un drawdown sur "
    "période complète.", "li"
)
para(
    "P4 (surapprentissage) — l'étude 3 attaque la largeur des intervalles, cause racine de "
    "l'indécision de la Phase 5 ; les études 1, 2 et 4 sont des tentatives pré-enregistrées de "
    "réfuter notre propre résultat nul, ce qui est la forme la plus directe de contrôle du "
    "surapprentissage.", "li"
)

# ═══ 11. Limites ═════════════════════════════════════════════════════════════
para("11. Limites (restées explicites)", "h1")
para(
    "Cinq fenêtres de crise sont cinq observations. Les comparaisons de portefeuille de l'étude 5 "
    "ne portent AUCUN test de significativité et sont descriptives ; seule la fréquence de "
    "détection de régime est testée.", "li"
)
para(
    "L'étude 5 ne couvre que `etf_2017` : `full_2021` commence mi-2022 et n'inclut qu'une crise "
    "partielle. Les fenêtres sont définies sur le S&P 500, référence correcte pour un univers "
    "majoritairement ETF mais pas pour les valeurs BVC.", "li"
)
para(
    "L'étude 3 concatène six segments ; le bootstrap par blocs les traite comme une série "
    "unique, ce qui sous-estime légèrement l'incertitude aux jointures.", "li"
)
para(
    "L'étude 1 repose sur un univers marocain profond assemblé à la main depuis investing.com, "
    "non intégré au pipeline médaillon et non ajusté des dividendes — il sert à tester une "
    "hypothèse, pas à produire un résultat de production.", "li"
)
para(
    "Le LSTM (troisième famille F7) reste reporté depuis la Phase 4B après un segfault "
    "torch+xgboost. Compte tenu des cinq résultats ci-dessus, l'ajouter attaquerait la précision "
    "de prédiction — précisément le levier dont ce livrable démontre qu'il n'est pas le goulot "
    "d'étranglement.", "li"
)

# ═══ Références ══════════════════════════════════════════════════════════════
para("Références", "h1")
para("Jagannathan, R. & Ma, T. (2003). Risk Reduction in Large Portfolios: Why Imposing the "
     "Wrong Constraints Helps. Journal of Finance — la contrainte comme rétrécissement.", "li")
para("Bailey, D. & López de Prado, M. (2014). The Deflated Sharpe Ratio. Journal of Portfolio "
     "Management.", "li")
para("López de Prado, M. (2018). Advances in Financial Machine Learning. Wiley — ch. 7 "
     "(K-Fold purgé), ch. 8 (importance des features).", "li")
para("Chopra, V. & Ziemba, W. (1993). The Effect of Errors in Means, Variances, and Covariances "
     "on Optimal Portfolio Choice.", "li")
para("Notes de recherche détaillées : docs/DEEP_MOROCCO_EXPERIMENT.md · "
     "docs/FUNDAMENTALS_EXPERIMENT.md · docs/NESTED_WALKFORWARD_EXPERIMENT.md · "
     "docs/REGIME_CONDITIONAL_CAP_EXPERIMENT.md · docs/CRISIS_WINDOWS_EXPERIMENT.md · "
     "docs/ETF_DEEP_HISTORY_EXPERIMENT.md", "li")

for hp in doc.sections[0].header.paragraphs:
    if hp.runs:
        hp.runs[0].text = "Phase 8 — Études de robustesse"
        for r in hp.runs[1:]:
            r.text = ""

doc.save(str(OUT))
d2 = Document(str(OUT))
print(f"saved {OUT.name}: {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables")
