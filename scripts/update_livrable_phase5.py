"""Bring docs/Livrable_Phase5_Evaluation_OOS.docx onto the forward-only protocol.

Addresses: P4 — this deliverable still described purged K-Fold as the ML
hyperparameter selector, and carried the frozen-test table produced under that
protocol. The validation-protocol phase replaced the selector with strictly
forward-only folds, which changed the selected hyperparameters and therefore
six of the table's eight numbers. A committed supervisor deliverable that
contradicts the governance package is exactly the drift Phase 3 exists to
remove — and CLAUDE.md §17.11 is the record of a claim surviving in six
surfaces at once because each was edited by hand.

WHY A SCRIPT AND NOT A HAND-EDIT. The Phase 8 builder clones THIS file as its
style template, so regenerating Phase 5 from scratch would risk changing the
house style of the whole set. Targeted edits keep the document identical
except where it is wrong, while the replacement text lives here as reviewable
source rather than only inside a binary. Numbers are read from `data/gold/`,
never typed.

Idempotent: every replacement is skipped when the paragraph already carries
the new wording, so re-running is a no-op.
"""

from __future__ import annotations

import copy
import json
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
GOLD = ROOT / "data" / "gold"
DOCX = ROOT / "docs" / "Livrable_Phase5_Evaluation_OOS.docx"

P5 = json.loads((GOLD / "phase5_results.json").read_text())
PAIRED = json.loads((GOLD / "paired_comparison_results.json").read_text())
PROTOCOL = json.loads((GOLD / "phase5_validation_protocol.json").read_text())

UNIVERSES = ("etf_2017", "full_2021")
FR_LABEL = {
    "rf_signal_tuned": "RF optimisé", "xgb_signal_tuned": "XGB optimisé",
    "regime_conditional": "regime_conditional", "equal_weight": "equal_weight",
}


def fr(x: float, n: int = 3) -> str:
    return f"{x:.{n}f}".replace(".", ",").replace("-", "−")


def _sharpe(universe: str, strategy: str) -> tuple[float, list[float]]:
    block = P5[universe]
    entry = block["tuned"].get(strategy) or block["baselines"][strategy]
    return entry["test_sharpe_net"], entry["test_sharpe_ci"]


def set_paragraph(paragraph, text: str, bold_lead: str | None = None) -> None:
    """Replace a paragraph's text, keeping its first run as the style carrier.

    Wiping and re-adding runs would drop the paragraph's font; reusing run[0]
    keeps it. `bold_lead` reproduces the document's convention of a bolded
    sentence opener.
    """
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    keeper = runs[0]
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    if bold_lead:
        keeper.text = bold_lead
        keeper.bold = True
        tail = copy.deepcopy(keeper._element)
        keeper._element.addnext(tail)
        paragraph.runs[1].text = text
        paragraph.runs[1].bold = False
    else:
        keeper.text = text
        keeper.bold = False


def replace_if_stale(paragraph, marker: str, text: str, bold_lead: str | None = None) -> bool:
    """Rewrite only when the paragraph still shows the superseded wording."""
    if marker not in paragraph.text:
        return False
    set_paragraph(paragraph, text, bold_lead)
    return True


def build_paired_section(doc, after_index: int) -> int:
    """Insert the paired-comparison section after §4, cloning local styles."""
    paragraphs = doc.paragraphs
    # Key on the heading, not on body prose: an earlier guard matched
    # "bootstrap par blocs pairé" while the inserted text writes "PAIRÉ", so the
    # section was appended again on every run.
    if any(p.text.strip().startswith("4 ter.") for p in paragraphs):
        return 0

    heading_src = next(p for p in paragraphs if p.style is not None
                       and p.style.name == "Heading 1" and p.text.startswith("4 bis."))
    body_src = paragraphs[after_index]
    anchor = paragraphs[after_index]._p

    blocks = [
        ("heading", "4 ter. Le test qui manquait — comparaison pairée des différences"),
        ("body",
         "La conclusion ci-dessus reposait sur des intervalles de confiance MARGINAUX, "
         "un par stratégie. Or le chevauchement de deux intervalles marginaux n'est pas "
         "un test de leur différence : les stratégies partagent les mêmes journées de "
         "marché, leurs erreurs sont corrélées, et deux intervalles peuvent se recouvrir "
         "largement alors que l'écart pairé est systématiquement d'un même signe. "
         "L'inverse vaut aussi : un non-chevauchement n'établirait pas davantage une "
         "différence."),
        ("body",
         "L'instrument manquant est un bootstrap par blocs PAIRÉ sur les mêmes dates de "
         "la fenêtre gelée, net de coûts : les deux séries sont rééchantillonnées avec "
         "les MÊMES indices de blocs, ce qui préserve la dépendance sérielle de chaque "
         "stratégie et la corrélation intra-journalière entre elles. La p-value est "
         "centrée sous l'hypothèse nulle — les différences rééchantillonnées sont "
         "recentrées sur zéro, puis on mesure la part de cette distribution au-delà de "
         "l'écart observé. Ce n'est pas la fraction brute de tirages positifs, rapportée "
         "séparément et jamais présentée comme une p-value."),
        ("body", _paired_verdict()),
        ("body",
         "Correction pour tests multiples : NON ÉTABLIE, et déclarée comme telle. Un "
         "test de White ou de Hansen exige la série de rendements sur la fenêtre gelée "
         "de CHAQUE configuration explorée ; seules les configurations finalement "
         "retenues en disposent. L'appliquer au sous-ensemble disponible sous-estimerait "
         "l'ampleur de la recherche et se lirait comme une fausse assurance."),
    ]

    inserted = 0
    for kind, text in blocks:
        source = heading_src if kind == "heading" else body_src
        clone = copy.deepcopy(source._p)
        anchor.addnext(clone)
        anchor = clone
        from docx.text.paragraph import Paragraph

        set_paragraph(Paragraph(clone, source._parent), text)
        inserted += 1
    return inserted


def _paired_verdict() -> str:
    comparisons = PAIRED["comparisons"]
    established = [c for c in comparisons
                   if c["sharpe_diff_ci"][0] > 0 and c["p_value_no_outperformance"] < 0.05]
    p_values = [c["p_value_no_outperformance"] for c in comparisons]
    highlight = next(
        c for c in comparisons
        if c["universe"] == "full_2021" and c["candidate"] == "xgb_signal_tuned"
        and c["benchmark"] == "regime_conditional"
    )
    xgb, _ = _sharpe("full_2021", "xgb_signal_tuned")
    regime, _ = _sharpe("full_2021", "regime_conditional")
    lo, hi = highlight["sharpe_diff_ci"]
    verdict = (
        "aucune n'établit de surperformance"
        if not established
        else f"{len(established)} établissent une surperformance"
    )
    return (
        f"Résultat : sur les {len(comparisons)} comparaisons pairées, {verdict} — "
        f"tous les intervalles contiennent zéro et les "
        f"p-values s'échelonnent de {fr(min(p_values))} à {fr(max(p_values))}. Le cas à "
        f"isoler est exactement celui que ce test sert à révéler : sur full_2021, le XGB "
        f"optimisé obtient {fr(xgb)} contre {fr(regime)} pour la stratégie à régimes et "
        f"l'artefact l'enregistre comme dépassant la haie, mais le test pairé donne un "
        f"écart de {fr(highlight['sharpe_diff'])}, IC 90 % [{fr(lo)} ; {fr(hi)}], "
        f"p = {fr(highlight['p_value_no_outperformance'])}. Le classement ponctuel n'est "
        f"pas soutenu par les données. Ce n'est pas non plus une démonstration "
        f"d'équivalence : ne pas rejeter n'est pas accepter l'hypothèse nulle."
    )


def update_table(doc) -> int:
    """Rewrite the frozen-test table from the current artifact."""
    table = doc.tables[0]
    changed = 0
    for row, universe in zip(table.rows[1:], UNIVERSES):
        block = P5[universe]
        label = (f"{universe} ({block['test_start'][:7].replace('-', '-')}"
                 f"→{block['test_end'][:7]})")
        cells = [label]
        for strategy in ("rf_signal_tuned", "xgb_signal_tuned",
                         "regime_conditional", "equal_weight"):
            value, (lo, hi) = _sharpe(universe, strategy)
            cells.append(f"{fr(value)} [{fr(lo, 2)} ; {fr(hi, 2)}]")
        for cell, text in zip(row.cells, cells):
            if cell.text.strip() != text:
                set_paragraph(cell.paragraphs[0], text)
                changed += 1
    return changed


def main() -> None:
    doc = Document(str(DOCX))
    paragraphs = doc.paragraphs
    edits = 0

    protocol_name = PROTOCOL["protocol"]
    config = PROTOCOL["config"]
    n_trials = P5["etf_2017"]["n_search_trials"]

    for paragraph in paragraphs:
        edits += replace_if_stale(
            paragraph, "Quand les intervalles se chevauchent, les stratégies sont",
            "Sharpe net du segment de test (fenêtre jamais vue par la sélection), avec "
            "intervalle de confiance à 90 % par bootstrap par blocs. Ces intervalles sont "
            "MARGINAUX : leur chevauchement ne teste pas la différence entre deux "
            "stratégies et n'autorise à conclure ni à une supériorité ni à une "
            "équivalence. Le test de la différence est en §4 ter.",
        )
        edits += replace_if_stale(
            paragraph, "3 h 21 de calcul : sélection par K-Fold purgé",
            "Deux corrections de données sont intervenues après la première rédaction de "
            "ce livrable. Contrairement aux livrables des phases 2, 4, 4B et 4C — annotés "
            "mais non recalculés — la Phase 5 a été INTÉGRALEMENT RÉ-EXÉCUTÉE sur les "
            "données corrigées le 27/07/2026, puis de nouveau après le remplacement du "
            "sélecteur par une validation strictement antérieure : choix des "
            "hyperparamètres, choix des leviers et test gelé refaits de bout en bout. Les "
            "chiffres de ce document sont ceux de cette dernière exécution.",
        )
        edits += replace_if_stale(
            paragraph, "Validation croisée purgée · sélection honnête",
            "Sélection strictement antérieure · comparaison pairée des différences · "
            "intervalles de confiance · Sharpe déflaté",
        )
        edits += replace_if_stale(
            paragraph, "K-Fold purgé et embargoé, noté par le coefficient",
            "sélection strictement ANTÉRIEURE (PurgedWalkForwardSplit : "
            f"{config.get('n_splits')} plis à origine glissante, embargo "
            f"{config.get('embargo_dates')} jours, horizon d'étiquette "
            f"{config.get('label_horizon')} jour), notée par le coefficient "
            "d'information (corrélation de rang de Spearman entre rendements prédits et "
            "réalisés : pour un signal, c'est l'ordre qui doit généraliser).",
            bold_lead="Hyperparamètres ML (profondeur/feuilles/taux des RF/XGB) — ",
        )
        edits += replace_if_stale(
            paragraph, "Le module purged_kfold.py (implémentation maison",
            "Le K-Fold purgé canonique de López de Prado autorise l'entraînement des "
            "deux CÔTÉS du pli de validation : défendable pour des étiquettes "
            "chevauchantes, mais il note un pli de 2018 avec un modèle partiellement "
            "ajusté sur 2019–2024, une question qu'aucun allocataire réel ne peut poser. "
            f"La sélection utilise donc {protocol_name} : chaque pli vérifie "
            "max(entraînement) < min(validation), affirmé au moment du découpage plutôt "
            "que supposé, et une histoire insuffisante lève une erreur au lieu de "
            "produire silencieusement moins de plis. PurgedKFold reste dans le dépôt et "
            "reste testé : la fonction n'a pas changé de sens, c'est l'appelant qui a "
            "changé d'outil.",
            bold_lead="Note anti-fuite (P4). ",
        )
        edits += replace_if_stale(
            paragraph, "sont statistiquement indiscernables sur les deux univers",
            "hors-échantillon, sur données gelées et avec de vraies barres d'erreur, "
            "AUCUNE comparaison pairée n'établit de surperformance des modèles F7 "
            "honnêtement optimisés sur la ligne de base regime_conditional, ni "
            "l'inverse, sur aucun des deux univers (§4 ter). L'estimation ponctuelle "
            "gagnante s'inverse même selon l'univers, soit l'inverse exact de "
            "l'évaluation précédente. Ce changement de signe obtenu en corrigeant les "
            "données et non en cherchant un résultat est la démonstration la plus nette "
            "que ces classements ponctuels relèvent du bruit. Les comparaisons "
            "apparemment tranchées des phases 4B–4C étaient dans le bruit depuis le "
            "début. Formulation à conserver : on ne dit ni « supérieur » ni "
            "« équivalent » sans test pairé.",
            bold_lead="Énoncée une fois pour qu'elle ne dérive pas : ",
        )
        edits += replace_if_stale(
            paragraph, "face à l'ensemble des 36 configurations",
            f"C'est un résultat fort, pas une déception. Le Sharpe déflaté de la "
            f"meilleure stratégie optimisée face à l'ensemble des {n_trials} "
            f"configurations de la recherche vaut "
            f"{fr(P5['etf_2017']['best_tuned_dsr_vs_search'], 3)} sur etf_2017 "
            f"(crédible) et {fr(P5['full_2021']['best_tuned_dsr_vs_search'], 3)} sur "
            f"full_2021 (le plus fragile des deux). Ce N a lui-même été corrigé : le "
            f"registre ne consignait que les essais de leviers et omettait la grille "
            f"d'hyperparamètres ML, sous-estimant l'ampleur de la recherche. Les "
            f"coefficients d'information de la validation croisée restent honnêtement "
            f"faibles, comme l'est habituellement la prédiction de rendements "
            f"financiers — ils ont pourtant à peu près doublé avec l'historique étendu, "
            f"sans que le portefeuille en profite : la quatrième confirmation du plafond "
            f"« précision de prédiction ≠ performance du portefeuille »."
        )
        edits += replace_if_stale(
            paragraph, "K-Fold purgé (sélection ML sans fuite), segment de test gelé",
            "sélection strictement antérieure (aucun pli n'entraîne après sa fenêtre de "
            "validation), segment de test gelé (verdict jamais vu par la sélection), "
            "bootstrap par blocs PAIRÉ sur les différences (le seul instrument qui teste "
            "une comparaison), intervalles de confiance, et un Sharpe déflaté qui compte "
            "toute la recherche — grille ML comprise, ce que le registre omettait.",
            bold_lead="P4 (surapprentissage du backtest) — le cœur de la phase : ",
        )
        edits += replace_if_stale(
            paragraph, "n'apporte aucun avantage statistiquement significatif",
            "La Phase 5 est l'aboutissement méthodologique du projet. Elle a construit "
            "la pile ML complète — détection de régime, covariance dynamique, signaux de "
            "prédiction adaptatifs, optimisation sensible aux coûts — puis l'a soumise à "
            "l'évaluation qu'exige le problème. Sa conclusion : aucune comparaison "
            "pairée n'établit d'avantage de la pile ML sur une ligne de base à "
            "commutation de régime bien construite, hors-échantillon. La machinerie "
            "d'évaluation qui le montre rigoureusement — et qui détecterait un vrai "
            "avantage si un futur modèle en produisait un — est le livrable. Une équipe "
            "qui rapporte un résultat nul sans fuite, avec barres d'erreur et un test de "
            "la différence, comprend le problème bien mieux qu'une équipe qui rapporte "
            "un Sharpe gonflé."
        )

    cells = update_table(doc)
    # Anchor AFTER §4 bis, not after §4: the nested walk-forward is dated
    # earlier and reads as a continuation of the verdict, so slotting the paired
    # test between them would put "4 ter" ahead of "4 bis".
    section_index = max(
        i for i, p in enumerate(doc.paragraphs)
        if "Conséquence sur le verdict de la section 3" in p.text
    )
    inserted = build_paired_section(doc, section_index)

    doc.save(str(DOCX))
    print(f"{DOCX.relative_to(ROOT)}: {edits} paragraph(s), {cells} table cell(s), "
          f"{inserted} new paragraph(s)")


if __name__ == "__main__":
    main()
