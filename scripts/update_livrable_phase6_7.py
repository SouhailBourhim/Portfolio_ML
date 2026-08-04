"""Remove the retracted significance wording from the Phase 6+7 deliverable.

Addresses: P4 — this committed, jury-facing deliverable still asserted that the
F7 signals add no "statistiquement significative" value and that the dashboard's
displayed gaps are "pas statistiquement significatifs". Both are the retracted
claim in its second form: failing to reject is not accepting the null, and
marginal intervals are not a test of a difference in EITHER direction.

Targeted edits rather than a rebuild: this file has no generator, and a
from-scratch reconstruction would risk its layout for two sentences. The
replacement text lives here as reviewable source instead of only inside a
binary — the same reasoning as `update_livrable_phase5.py`.

Idempotent: each replacement is skipped once the new wording is present.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "Livrable_Phase6-7_Suite_Portfolio_ML.docx"

REPLACEMENTS = (
    (
        "qu'ils n'ajoutent pas de valeur statistiquement significative",
        "qu'aucun avantage ne leur est démontrable : les comparaisons pairées "
        "menées depuis n'établissent de surperformance dans aucun sens",
    ),
    (
        "avec un avertissement rappelant que les écarts ne sont pas "
        "statistiquement significatifs",
        "avec un avertissement rappelant qu'aucun test pairé n'établit ces "
        "écarts, et que des intervalles marginaux ne testent pas une différence",
    ),
)


def set_paragraph(paragraph, text: str) -> None:
    """Rewrite a paragraph, keeping run[0] as the style carrier."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    keeper = runs[0]
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    keeper.text = text


def main() -> None:
    document = Document(str(DOCX))
    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text
        updated = text
        for old, new in REPLACEMENTS:
            if old in updated:
                updated = updated.replace(old, new)
        if updated != text:
            set_paragraph(paragraph, updated)
            changed += 1
    document.save(str(DOCX))
    print(f"{DOCX.relative_to(ROOT)}: {changed} paragraph(s) rewritten")


if __name__ == "__main__":
    main()
